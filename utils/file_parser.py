import PyPDF2
from docx import Document
import re

def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text
    except Exception as e:
        print(f"Error reading PDF {file_path}: {str(e)}")
        return ""

def extract_text_from_doc(file_path):
    """Extract text from DOCX file."""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text
        return text
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {str(e)}")
        return ""

def clean_text(text):
    """Clean and normalize text."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    return text.strip()

def normalize_resume_text(text):
    """Normalize resume text for processing."""
    text = clean_text(text)
    return text
